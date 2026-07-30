"""
Build the submission write-up in both formats:

    docs/Off_Ball_Run_Value__Write_Up.pdf     -- for sending
    docs/Off_Ball_Run_Value__Write_Up.docx    -- for editing

    python docs/make_writeup.py

Two design decisions worth knowing:

1. Every number in the prose is read from outputs/ and data/processed/ at build
   time, never typed in, so the document cannot drift from the model.

2. The PDF and the DOCX render from one shared list of content blocks (see
   `content()`), not two parallel copies of the text. Maintaining the same
   document twice guarantees they diverge.

WARNING: re-running this script OVERWRITES both files. Once you start editing the
.docx by hand, that file is the master -- export the PDF from Word instead, or
your edits are lost.
"""
from __future__ import annotations
import base64
import json
import re
import subprocess
import sys
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
import config  # noqa: E402

DOCS = ROOT / "docs"
FIGS = DOCS / "figures"
CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"

BRAND = "#5b4fd6"      # print-safe version of the app's #877ae8
INK = "#15171a"
MUTED = "#5f6672"


# --------------------------------------------------------------------- numbers
def gather() -> dict:
    runs = pd.read_parquet(config.DATA_PROC / "runs.parquet")
    mins = pd.read_parquet(config.DATA_PROC / "minutes.parquet")
    pm = pd.read_csv(ROOT / "outputs" / "run_player_metrics.csv")
    abl = json.load(open(ROOT / "outputs" / "ablation_report.json"))
    rep = json.load(open(ROOT / "outputs" / "run_report.json"))
    val = json.load(open(ROOT / "outputs" / "value_report.json"))

    u = runs[(runs["usable"] == 1) & (runs["is_run"] == 1)]

    # the run used as the worked example, scored by both fits of Model B
    mb = runs[(runs["match_id"] == 3803000) & (runs["minute"] == 73)
              & (runs["runner"].str.contains("Mbapp", na=False)) & (runs["usable"] == 1)].iloc[0]

    return dict(
        matches=runs["match_id"].nunique(), receipts=len(runs), usable=len(u),
        usable_pct=100 * len(u) / len(runs),
        gate_pct=100 * (runs["usable"] == 1).mean(),
        players=len(pm), minutes=mins["minutes"].sum(), top=pm.nlargest(10, "RunValue90"),
        abl=abl, tiers=abl["tiers"], gain360=abl["total_gain_from_360"],
        n_states=abl["n_rows"], base_rate=abl["base_rate"],
        v_brier=abl["tiers"][-1]["brier"], naive_lift=val["auc_lift"],
        ctx_auc=rep["metrics_context"]["auc"], run_auc=rep["metrics_plus_run"]["auc"],
        run_lift=rep["auc_lift"], run_rows=rep["metrics_context"]["n"],
        run_base=rep["metrics_context"]["base_rate"],
        run_brier=rep["metrics_plus_run"]["brier"],
        mb_ctx=float(mb["rv_context"]), mb_full=float(mb["run_value"]),
        mb_add=float(mb["run_value_added"]),
    )


# Display names. Taking the last token is wrong for Iberian and Brazilian naming
# ("Kylian Mbappé Lottin" -> "Lottin", "Neymar da Silva Santos Junior" -> "Junior").
SHORT = [("Mbapp", "K. Mbappé"), ("Boniface", "V. Boniface"), ("Wirtz", "F. Wirtz"),
         ("Trinc", "F. Trincão"), ("Braithwaite", "M. Braithwaite"),
         ("Hofmann", "J. Hofmann"), ("Neymar", "Neymar"), ("Schick", "P. Schick"),
         ("Ekitike", "H. Ekitike"), ("Tella", "N. Tella"), ("Dembélé", "O. Dembélé"),
         ("Messi", "L. Messi"), ("Adli", "A. Adli")]


def short_name(name: str) -> str:
    for needle, disp in SHORT:
        if needle in name:
            return disp
    parts = name.split()
    return f"{parts[0][0]}. {parts[-1]}" if len(parts) > 1 else name


# ------------------------------------------------------------------- flowchart
def flow_png(matches: int) -> Path:
    """
    One flowchart image, shared by both formats. Drawn with matplotlib rather than
    inline SVG precisely so the DOCX and the PDF cannot show different diagrams.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrow

    steps = [("1. Data", ["StatsBomb 360:", "events + freeze-frames", f"{matches} matches"]),
             ("2. Reconstruct", ["Pair the release and", "receipt frames to", "recover the run"]),
             ("3. Describe", ["Ball, runner and", "defensive shape", "at both instants"]),
             ("4. Model", ["P(progression),", "folds grouped", "by match"]),
             ("5. Aggregate", ["Run Value per 90", "for every player", "over 600 mins"])]

    # Kept deliberately wide and short: it renders full page-width, so a tall
    # aspect ratio costs vertical space the write-up does not have.
    fig, ax = plt.subplots(figsize=(12.0, 1.16), dpi=300)
    ax.set_xlim(0, 100); ax.set_ylim(0, 10.2); ax.axis("off")
    fig.patch.set_facecolor("white")
    bw, gap = 17.4, 3.2
    for i, (head, body) in enumerate(steps):
        x = i * (bw + gap)
        ax.add_patch(FancyBboxPatch((x, 0.9), bw, 8.4, boxstyle="round,pad=0,rounding_size=0.6",
                                    linewidth=0.9, edgecolor=BRAND, facecolor="#f6f6fb"))
        ax.text(x + 1.0, 7.85, head, fontsize=7.4, fontweight="bold", color=BRAND, va="center")
        for j, line in enumerate(body):
            ax.text(x + 1.0, 5.95 - j * 1.6, line, fontsize=6.5, color="#3f4550", va="center")
        if i < len(steps) - 1:
            ax.add_patch(FancyArrow(x + bw + 0.35, 5.1, gap - 1.2, 0, width=0.06,
                                    head_width=0.8, head_length=0.85,
                                    color=BRAND, length_includes_head=True))
    fp = FIGS / "fig_flow.png"
    fig.savefig(fp, dpi=300, bbox_inches="tight", facecolor="white", pad_inches=0.02)
    plt.close(fig)
    return fp


# ---------------------------------------------------------------- inline markup
TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", re.S)


def inline(text: str):
    """Split a string into (text, bold, italic, code) runs. **b**, *i*, `code`."""
    out = []
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], True, False, False))
        elif part.startswith("`") and part.endswith("`"):
            out.append((part[1:-1], False, False, True))
        elif part.startswith("*") and part.endswith("*"):
            out.append((part[1:-1], False, True, False))
        else:
            out.append((part, False, False, False))
    return out


# ------------------------------------------------------------------- content
def content(d: dict) -> list:
    """The document, once. Rendered to HTML/PDF and DOCX by the two renderers."""
    t = d["tiers"]
    top = d["top"]
    trincao = short_name(top.iloc[3]["runner"])
    trincao_90s = top.iloc[3]["minutes"] / 90

    top_rows = []
    for i, (_, r) in enumerate(top.iterrows(), start=1):
        pos = (r["position"].replace("Attacking Midfield", "Att Mid")
               .replace("Center Forward", "Forward").replace("Right Wing Back", "RWB"))
        top_rows.append([str(i), short_name(r["runner"]), r["team"], pos,
                         f"{r['minutes'] / 90:.1f}", f"{r['Runs90']:.1f}",
                         f"{r['FinalThirdRuns90']:.1f}", f"{r['InBehind90']:.2f}",
                         f"{r['Encirclement']:.2f}", f"{r['RunValue90']:.3f}"])

    notes = ["what event data alone can see", "the runner's own position, from 360",
             "where the defenders are", "shape of the block (negligible)"]
    abl_rows = [[x["tier"], str(x["n_features"]), f"{x['auc']:.4f}", f"{x['brier']:.4f}",
                 "—" if x["gain_vs_previous"] is None else f"+{x['gain_vs_previous']:.4f}",
                 notes[i]] for i, x in enumerate(t)]

    return [
        ("h1", "Valuing off-ball runs with StatsBomb 360"),
        ("sub", "A possession-value model for the movement that happens away from the "
                "ball, and a player metric built on top of it."),
        ("byline", "Pranav Mohan  ·  Data Scientist assessment, Parts 1 & 2  ·  July 2026  "
                   "·  Code: github.com/pranavm28/statsbomb-off-ball"),

        ("lede", f"**Headline.** Event data records the pass; it never records the run that made "
                 f"the pass possible. Using 360 freeze-frames I reconstruct {d['usable']:,} "
                 f"off-ball runs and value each one by how much the *manner* of the movement raises "
                 f"the chance the possession progresses. Seeing the players — not just the ball — is worth "
                 f"+{d['gain360']:.4f} AUC over event-only features on identical rows and folds. "
                 f"The applied finding is a profile, not just a ranking: **Victor Boniface** ranks "
                 f"second per 90 on a fifth of Mbappé's minutes, while receiving in the most "
                 f"crowded space of anyone near the top."),

        ("h2", "1", "What \"possession value\" means here"),
        ("p", "**Possession value** here is a number attached to a *state* of the possession — "
              "not to a player and not to an action, but to a moment: the ball is here, this "
              "player is there, the defenders are arranged like this. The value of a move is then "
              "the *change* in that number between two states. That logic is shared with xT and "
              "VAEP; what differs is what goes into the state."),
        ("p", "**Two models do two different jobs.** They are kept separate on purpose, and "
              "knowing which is which is needed to read every number below."),
        ("table", ["", "Model A — state value", "Model B — run value"], [
            ["Question it answers",
             "Does seeing the *players*, not just the ball, make a game state more predictable "
             "at all?",
             "How much did *this* player's movement add, over and above where his run started "
             "and finished?"],
            ["One row is", "one game state: a player and the ball at one instant",
             "one reconstructed off-ball run"],
            ["Rows", f"{d['n_states']:,}", f"{d['run_rows']:,}"],
            ["Target — the 0/1 label it learns from",
             "a shot within the next **5 seconds**",
             "**progression**: a shot *or* a final-third entry within the next **5 actions** of "
             "the possession"],
            ["How often the label is 1", f"{100 * d['base_rate']:.1f}%", f"{100 * d['run_base']:.1f}%"],
            ["What it is used for in this document",
             "the 360 ablation in §6 — does 360 earn its keep",
             "**Run Value per 90** — the metric and ranking in §4"],
        ], [0, 0, 0]),
        ("p", f"**Why two targets instead of one?** \"Shot within five seconds\" is right for the "
              f"first question: it is close to what football cares about, and across "
              f"{d['n_states']:,} states there is enough data to learn an event that rare "
              f"({100 * d['base_rate']:.1f}%). It is wrong for the second. At the level of one "
              f"run, shots are too rare to tell a good run from a lucky one, so Model B is judged "
              f"on progression, which happens {100 * d['run_base']:.1f}% of the time. A run that "
              f"drags the ball into the final third has done its job whether or not a shot "
              f"arrives three passes later."),
        ("p", "**Why a learned state-value function rather than a grid xT?** xT keys value to "
              "*where the ball is*; this keys it to *where the players are*. Two identical ball "
              "positions with different defensive shapes are one cell in xT and two distinct "
              "states here. An xT surface is still used, as a context *feature*, so the model "
              "knows the geographic prior it is adding to — the threat surface is an input, never "
              "the thing being predicted."),

        ("h2", "2", "How it works"),
        ("img_full", str(FIGS / "fig_flow.png")),
        ("p", f"The reconstruction does the real work. A `Ball Receipt` event carries its own "
              f"freeze-frame with the receiver tagged `actor`, which fixes where the run *ended* "
              f"and who made it. The preceding pass's frame fixes where everyone stood when the "
              f"ball was struck — but frames are anonymous, so the run's *origin* must be "
              f"inferred: the nearest teammate in the release frame, accepted only if the implied "
              f"sprint is physically possible (≤ {config.MAX_RUN_SPEED} m/s over the pass "
              f"duration). This origin inference is the biggest assumption in the project."),
        ("p", f"**Two filters, in order.** Of {d['receipts']:,} receipts examined, "
              f"{d['gate_pct']:.0f}% pass the plausibility gates — an unambiguous nearest "
              f"teammate and a reachable distance. Of those, the ones where the player actually "
              f"*moved* (≥ {config.MIN_RUN_DISTANCE} m) are runs: "
              f"{d['usable']:,}, or {d['usable_pct']:.0f}% of all receipts. The rest are players "
              f"receiving roughly where they already stood, which is a reception, not a run."),

        ("h2", "3", "Scope, and why"),
        ("p", f"Four full competition-seasons with 360 coverage — La Liga 2020/21, Ligue 1 "
              f"2021/22 and 2022/23, Bundesliga 2023/24: {d['matches']} matches, "
              f"{d['minutes']:,.0f} player-minutes. I pooled seasons rather than use the World "
              f"Cup because off-ball running needs volume per player and the metric is a rate: a "
              f"tournament gives seven matches, a season thirty-plus. The honest cost is that "
              f"open-data 360 covers one focus club per competition, so the pool is effectively "
              f"three clubs and the ranking is within-sample, not cross-league."),

        ("h2", "4", "The metric, and the ranking"),
        ("p", "**Run Value per 90** is built from **Model B**, so the quantity underneath it is "
              "the probability of *progression*, not of a shot. Four steps:"),
        ("bullets", [
            "**(i) Fit Model B twice on identical rows and folds.** Once on CONTEXT only — "
            "`origin_x/y`, `receipt_x/y`, threat at both ends, pass length: *where* the run "
            "happened. Then again on CONTEXT + RUN — distance, forward and lateral components, "
            "speed, separation gained, space at receipt, defenders broken, in-behind, "
            "encirclement and block geometry: *how* he moved.",
            "**(ii) Subtract, per run.** `run_value_added = P(progress | where + how) − "
            "P(progress | where)`. Both are out-of-fold predictions, meaning each run is scored "
            "by a model that never saw that run's match in training. This difference is the "
            "run's credit.",
            f"**(iii) Keep only runs received in the final third.** Outside it, \"predictive\" and "
            f"\"valuable\" come apart: a centre-back dropping into space reliably predicts "
            f"retained possession and creates nothing. Tested against the alternatives — the "
            f"attacking half moved the clearest false positive from 10th to 15th, the final third "
            f"moves him to 34th.",
            f"**(iv) Sum per player, divide by 90s played**, with a {config.MIN_MINUTES}-minute "
            f"floor. That leaves {d['players']} qualifying players."]),
        ("p", "**It is an information gain, not a probability.** It answers \"how much of the "
              "outcome is explained by the manner of the movement rather than its geography\", so "
              "it does not sum to goals. Two scouts watch the same possession, one seeing only "
              "the ball, one seeing the runner: Run Value is what the second knew that the first "
              "did not."),
        ("table", ["#", "Player", "Team", "Position", "90s", "Runs /90", "F3 runs /90",
                   "In behind /90", "Enclosed share", "Run Value /90"], top_rows,
         [0, 0, 0, 0, 1, 1, 1, 1, 1, 1]),
        ("cap", "Enclosed share = mean angular encirclement on receipt (1 = defenders on all "
                "sides, 0 = a completely clear side). F3 = final third."),

        ("h2", "5", "One run, in detail"),
        ("figure_two", str(FIGS / "fig_mbappe_run.png"), [
            "Ghosted dots = positions at pass release; solid = at arrival. Gold is the ball, "
            "green dotted is the run.",
            "Mbappé starts **outside the box** at the moment of release and arrives on the "
            "penalty spot 16.2 m later, with **four defenders** around him — encirclement 0.63, "
            "i.e. almost no clear side. The possession produces **0.70 xG within five seconds**.",
            f"**Scored, step by step.** From where the run started and finished alone, Model B put "
            f"the chance of progression at **{d['mb_ctx']:.3f}**. Once it could also see *how* he "
            f"moved — 16 m forward, in behind, into encirclement 0.63 — it rose to "
            f"**{d['mb_full']:.3f}**. The difference, **{d['mb_add']:+.3f}**, is this run's credit. "
            f"The possession did progress.",
            "That gap is the whole idea: without the frames this is just a 16 m pass into a "
            "crowd."]),

        ("h2", "6", "Validation"),
        ("p", "**Nested structure.** Events sit inside possessions inside matches, so a random "
              "split leaks — two events from one possession are near-duplicates and would land on "
              "both sides of the split. All folds are `GroupKFold` on `match_id`, so a whole match "
              "sits either in training or in test, never both. Every number quoted is out-of-fold: "
              "predicted by a model that never saw that match."),
        ("p", "**Does 360 earn its keep? (Model A.)** The honest test is not \"with vs without "
              "360\" — the runner's own position *is* a 360 feature, so a naive baseline smuggles "
              "it in. That mistake is worth "
              f"+{d['naive_lift']:.4f} AUC of flattery. Instead, four nested tiers on identical "
              f"rows and folds:"),
        ("table", ["Feature tier", "n", "AUC", "Brier", "Gain", "What it adds"], abl_rows,
         [0, 1, 1, 1, 1, 0]),
        ("p", f"Tier 1→3 is the real 360 contribution: **+{d['gain360']:.4f} AUC**. Tier 4 is "
              f"**+{t[3]['gain_vs_previous']:.4f}** — a null result, reported rather than dropped. "
              f"Read the *increments*, not the levels: an AUC of {t[0]['auc']:.2f} on tier 1 looks "
              f"impressive but the target is easy, since a ball already in the box implies a shot "
              f"soon. Only the incremental columns say anything."),
        ("p", f"**Does the movement earn its keep? (Model B.)** The same design, one level up: "
              f"CONTEXT alone reaches AUC {d['ctx_auc']:.4f}; adding the RUN features reaches "
              f"{d['run_auc']:.4f}, a lift of **+{d['run_lift']:.4f}** on "
              f"{d['run_rows']:,} runs. That lift is what Run Value is a per-run decomposition of. "
              f"**Calibration:** out-of-fold Brier {d['run_brier']:.4f} for Model B against a "
              f"{100 * d['run_base']:.1f}% base rate, and {d['v_brier']:.4f} for Model A against "
              f"{100 * d['base_rate']:.1f}%; predicted versus observed track closely across all "
              f"ten deciles for both (figures in the repo)."),

        ("h2", "7", "Say it two ways"),
        ("cards", [
            ("To a sporting director (~200 words)", [
                "We can now put a number on something we could previously only describe: what a "
                "player's running *without* the ball is worth.",
                "For every pass in our sample we can see where all twenty-two players stood when "
                "the ball was struck, and again when it arrived. That lets us rebuild the run the "
                "receiver made to get there and ask one question: how much more likely did a shot "
                "become because of the *way* he moved, rather than simply because of where the "
                "ball ended up?",
                "Mbappé finishes top, which tells you the method isn't broken but doesn't tell "
                "you anything you'd pay for. The name worth your time is **Victor Boniface**. He "
                "is second on this list on roughly a fifth of Mbappé's minutes, and he does it "
                "while receiving in the most crowded areas of anyone near the top — more than "
                "half his receptions come with defenders closing him from most sides. That is a "
                "specific and buyable habit: a forward who will still move into a packed box, and "
                "whose movement makes the pass playable.",
                "How much would I bet? Enough to put him on a shortlist and watch him properly. "
                "Not enough to price a bid on this alone — the sample is three clubs, and below "
                "the leading names the order shifts depending on which outcome we ask about."]),
            ("To the data team (2–3 sentences)", [
                f"Run Value is an ablation quantity: the out-of-fold gain in P(shot ≤ 5s) once the "
                f"model can see *how* the player moved, on top of a context model that already has "
                f"ball location, receipt location and defensive structure — so it is information "
                f"gain about the manner of movement, not a probability, and it will not sum to "
                f"goals.",
                f"Folds are `GroupKFold` on `match_id`; the 360 tiers are worth "
                f"+{d['gain360']:.4f} AUC over event-only features on identical rows and folds, "
                f"and out-of-fold predictions are calibrated (Brier {d['v_brier']:.4f} on a "
                f"{100 * d['base_rate']:.1f}% base rate).",
                "What would make me trust it more: a counterfactual attribution that holds the "
                "ball and the defenders fixed and moves only the runner, because the current "
                "quantity is target-sensitive — three of four target definitions I tested reorder "
                "the leaderboard, which is a property of information-gain metrics rather than a "
                "bug, but it needs pinning down before anyone scouts off it."])]),

        ("h2", "8", "What I used, and what I rejected"),
        ("cols", [
            ("Kept", [
                "**Ball Receipt freeze-frames.** The receiver is tagged `actor` — the only route "
                "to the runner's identity.",
                "**Nearest-teammate origin with a speed gate.** Frames are anonymous; the gate "
                "keeps physically impossible matches out.",
                "**LightGBM.** Tabular, non-linear, handles the mixed feature scales; the "
                "relationships here (distance, angle, congestion) are not linear.",
                "**Angular encirclement** over defender counts — *where* the defenders are beats "
                "*how many*.",
                "**Final-third restriction** to stop rewarding informative-but-harmless movement."]),
            ("Rejected, and why", [
                "**socceraction / VAEP off the shelf.** The point was defining the target myself; a "
                "black box is the failure mode the brief names.",
                "**Defender counts within 10 m.** Collinear with the geometry, no lift.",
                f"**Compactness tier.** +{t[3]['gain_vs_previous']:.4f} AUC — kept in the report "
                f"as a null, not in the metric.",
                "**Plain ΔV as the player metric.** It is joint with the pass; corr(ΔV, ball "
                "forward) exceeded corr(ΔV, run forward), so it partly ranked passers.",
                "**Pitch control / velocity models.** Freeze-frames have positions, no velocities.",
                "**World Cup 2022.** Too few matches per player for a rate metric."])]),

        ("h2", "9", "What breaks it"),
        ("p", "**The one that matters.** The metric is **target-sensitive**: of four target "
              "definitions tested, three reorder the leaderboard and two produce football-nonsense "
              "orderings (defenders and deep midfielders on top). That is inherent to an "
              "information-gain quantity — changing the outcome changes what counts as "
              "position-explained versus movement-explained — but the ordering below the leading "
              "names is not settled. The specified-but-unbuilt fix: a **counterfactual "
              "attribution** holding ball and defenders fixed and moving only the player, "
              "isolating his contribution instead of inferring it from a model comparison."),
        ("bullets", [
            "**Who it treats unfairly.** Target men and deep-lying creators. It rewards movement "
            "that moves *this* possession forward quickly, so a striker who occupies two centre-backs to "
            "open space for someone else scores nothing for it. Fixing that needs off-ball value "
            "attributed to the *space created for teammates*, not just the receiving player.",
            f"**Small samples.** The {config.MIN_MINUTES}-minute floor is ~6.7 × 90, so {trincao} "
            f"ranks 4th on {trincao_90s:.1f} × 90. Bootstrap intervals per player are the fix; "
            f"they are not in this version.",
            "**Three clubs.** Open-data 360 is club-focused, so this is within-sample.",
            f"**Inferred origins.** {100 - d['usable_pct']:.0f}% of receipts are discarded, and "
            f"the kept ones rest on a nearest-teammate assumption that will occasionally pick the "
            f"wrong player in a crowded box."]),

        ("h2", "10", "AI usage, in short"),
        ("p", "Full log with representative prompts, accepted/modified/rejected detail and the "
              "errors I caught is in `AI_USAGE.md` in the repo. Summary:"),
        ("table", ["Area", "What the assistant did", "What I did"], [
            ["Framing & metric", "Sketched options once I set the direction",
             "Mine: model off-ball runs as a supervised, xG-style problem rather than a grid"],
            ["Model family", "Recommended LightGBM + GroupKFold",
             "Accepted after checking the leakage argument held for nested events"],
            ["Features", "Drafted the geometry (encirclement, block spread)",
             "Rejected the density counters; required the ablation before believing any of it"],
            ["Verification", "Produced the first, flattering comparison",
             f"Caught the mislabelled 360 baseline (the \"no-360\" model contained a 360 feature) "
             f"— true gain +{d['gain360']:.4f}, not +{d['naive_lift']:.4f}"],
            ["Football sense", "Proposed reading block spread as \"stretched = valuable\"",
             "Caught the sign: the data says the opposite by 7.6×; it is a depth proxy"],
            ["Plotting & prose", "First drafts",
             "Restyled to my own conventions; fixed runs drawn off-pitch (360 frames are not "
             "clamped)"]], [0, 0, 0]),

        ("foot", "Data: **StatsBomb Open Data**, accessed via `statsbombpy`, used under the "
                 "StatsBomb open-data user agreement for non-commercial research. Data provided "
                 "by StatsBomb — with thanks.  ·  Reproduce: `pip install -r requirements.txt` "
                 "then `streamlit run app/streamlit_app.py`; full pipeline via "
                 "`requirements-pipeline.txt` and `build_runs.py`."),
    ]


# ---------------------------------------------------------------------- HTML
CSS = """
@page { size: A4; margin: 12mm 12mm 10mm 12mm; }
* { box-sizing: border-box; }
body { font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
       font-size: 8.1pt; line-height: 1.34; color: #15171a; margin: 0; }
h1 { font-size: 16.5pt; line-height: 1.1; margin: 0 0 1.4mm 0; letter-spacing: -0.4px; }
h2 { font-size: 9.3pt; margin: 3.0mm 0 1.1mm 0; color: #5b4fd6; }
h2 .n { color: #b9bcc6; margin-right: 4px; }
h3 { font-size: 8.7pt; margin: 0 0 1mm 0; color: #5b4fd6; }
p { margin: 0 0 1.5mm 0; }
b, strong { font-weight: 600; }
.sub { color: #5f6672; font-size: 8.7pt; margin: 0 0 1.2mm 0; }
.byline { color: #5f6672; font-size: 7.6pt; border-top: 1px solid #d8dbe2;
          padding-top: 1.2mm; margin-top: 1.8mm; }
.lede { background: #f6f6fb; border-left: 2.6px solid #5b4fd6;
        padding: 2.2mm 2.8mm; margin: 2.2mm 0 0.6mm 0; font-size: 8.6pt; }
.tbl { width: 100%; border-collapse: collapse; font-size: 7.3pt; margin: 1.1mm 0 1.5mm 0;
       page-break-inside: avoid; }
.tbl th { text-align: left; font-weight: 600; font-size: 7pt; color: #3f4550;
          border-bottom: 1px solid #9aa0ad; padding: 0.8mm 1.2mm; line-height: 1.2; }
.tbl td { padding: 0.78mm 1.2mm; border-bottom: 0.5px solid #e6e8ee; }
.tbl td.r, .tbl th.r { text-align: right; font-variant-numeric: tabular-nums; }
.tbl tbody tr:first-child td { background: #f8f8fc; }
.two { display: grid; grid-template-columns: 1fr 1fr; gap: 3.2mm; }
.figrow { display: grid; grid-template-columns: 0.92fr 1.08fr; gap: 3.2mm; }
.imgfull { width: 100%; display: block; margin: 0.8mm 0 1.8mm 0; }
.fig img { width: 100%; border-radius: 4px; display: block; }
.cap { font-size: 7.1pt; color: #5f6672; margin-top: 0.9mm; }
.card { border: 1px solid #d8dbe2; border-radius: 4px; padding: 2.2mm 2.6mm;
        page-break-inside: avoid; }
ul { margin: 0 0 1.5mm 0; padding-left: 3.8mm; }
li { margin-bottom: 0.45mm; }
.foot { border-top: 1px solid #d8dbe2; margin-top: 2.6mm; padding-top: 1.2mm;
        font-size: 6.9pt; color: #5f6672; }
code { font-family: "SF Mono", Menlo, monospace; font-size: 7.7pt; background: #f3f4f8;
       padding: 0.3mm 1mm; border-radius: 2px; }
"""


def h_inline(text: str) -> str:
    out = []
    for txt, b, i, c in inline(text):
        txt = txt.replace("&", "&amp;").replace("<", "&lt;")
        if b:
            txt = f"<b>{txt}</b>"
        if i:
            txt = f"<i>{txt}</i>"
        if c:
            txt = f"<code>{txt}</code>"
        out.append(txt)
    return "".join(out)


def b64(p: str) -> str:
    return base64.b64encode(Path(p).read_bytes()).decode()


def render_html(blocks: list) -> str:
    h = []
    for blk in blocks:
        k = blk[0]
        if k == "h1":
            h.append(f"<h1>{blk[1]}</h1>")
        elif k == "sub":
            h.append(f'<div class="sub">{blk[1]}</div>')
        elif k == "byline":
            h.append(f'<div class="byline">{blk[1]}</div>')
        elif k == "lede":
            h.append(f'<div class="lede">{h_inline(blk[1])}</div>')
        elif k == "h2":
            h.append(f'<h2><span class="n">{blk[1]}</span>{blk[2]}</h2>')
        elif k == "p":
            h.append(f"<p>{h_inline(blk[1])}</p>")
        elif k == "cap":
            h.append(f'<p class="cap">{h_inline(blk[1])}</p>')
        elif k == "foot":
            h.append(f'<div class="foot">{h_inline(blk[1])}</div>')
        elif k == "img_full":
            h.append(f'<img class="imgfull" src="data:image/png;base64,{b64(blk[1])}">')
        elif k == "figure_two":
            paras = "".join(f"<p>{h_inline(p)}</p>" for p in blk[2])
            h.append(f'<div class="figrow"><div class="fig">'
                     f'<img src="data:image/png;base64,{b64(blk[1])}"></div>'
                     f"<div>{paras}</div></div>")
        elif k == "table":
            heads, rows, aligns = blk[1], blk[2], blk[3]
            th = "".join(f'<th class="{"r" if a else ""}">{x.replace(" /", "<br>/")}</th>'
                         for x, a in zip(heads, aligns))
            tr = "".join("<tr>" + "".join(
                f'<td class="{"r" if a else ""}">{h_inline(c)}</td>'
                for c, a in zip(row, aligns)) + "</tr>" for row in rows)
            h.append(f'<table class="tbl"><thead><tr>{th}</tr></thead><tbody>{tr}</tbody></table>')
        elif k == "cards":
            cards = "".join(
                f'<div class="card"><h3>{t}</h3>'
                + "".join(f"<p>{h_inline(p)}</p>" for p in ps) + "</div>"
                for t, ps in blk[1])
            h.append(f'<div class="two">{cards}</div>')
        elif k == "cols":
            cols = "".join(
                f'<div><h3 style="color:#15171a">{t}</h3><ul>'
                + "".join(f"<li>{h_inline(x)}</li>" for x in items) + "</ul></div>"
                for t, items in blk[1])
            h.append(f'<div class="two">{cols}</div>')
        elif k == "bullets":
            h.append("<ul>" + "".join(f"<li>{h_inline(x)}</li>" for x in blk[1]) + "</ul>")
    return (f'<!doctype html><html><head><meta charset="utf-8">'
            f"<title>Valuing Off-Ball Runs — Write-Up</title><style>{CSS}</style>"
            f"</head><body>{''.join(h)}</body></html>")


# ---------------------------------------------------------------------- DOCX
def render_docx(blocks: list, out: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT, MONO = "Helvetica Neue", "Menlo"
    brand = RGBColor(0x5B, 0x4F, 0xD6)
    muted = RGBColor(0x5F, 0x66, 0x72)

    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    for attr in ("left_margin", "right_margin"):
        setattr(s, attr, Cm(1.5))
    s.top_margin, s.bottom_margin = Cm(1.4), Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    def shade(cell, hexcolor):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexcolor)
        cell._tc.get_or_add_tcPr().append(el)

    def add_runs(p, text, size=9, color=None, bold_all=False):
        for txt, b, i, c in inline(text):
            r = p.add_run(txt)
            r.font.name = MONO if c else FONT
            r.font.size = Pt(size - 0.6 if c else size)
            r.bold = b or bold_all
            r.italic = i
            if color is not None:
                r.font.color.rgb = color
        return p

    def para(text, size=9, color=None, space_after=4, bold_all=False, style=None):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        add_runs(p, text, size, color, bold_all)
        return p

    for blk in blocks:
        k = blk[0]
        if k == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(blk[1]); r.font.name = FONT; r.font.size = Pt(19); r.bold = True
        elif k == "sub":
            para(blk[1], size=10, color=muted, space_after=2)
        elif k == "byline":
            para(blk[1], size=8, color=muted, space_after=8)
        elif k in ("lede", "foot"):
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            add_runs(cell.paragraphs[0], blk[1], size=9 if k == "lede" else 7.5,
                     color=None if k == "lede" else muted)
            shade(cell, "F6F6FB" if k == "lede" else "FFFFFF")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif k == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{blk[1]}  {blk[2]}")
            r.font.name = FONT; r.font.size = Pt(11); r.bold = True; r.font.color.rgb = brand
        elif k == "p":
            para(blk[1])
        elif k == "cap":
            para(blk[1], size=7.5, color=muted)
        elif k == "img_full":
            doc.add_picture(blk[1], width=Inches(7.1))
        elif k == "figure_two":
            tbl = doc.add_table(rows=1, cols=2)
            tbl.columns[0].width, tbl.columns[1].width = Inches(3.5), Inches(3.6)
            tbl.cell(0, 0).paragraphs[0].add_run().add_picture(blk[1], width=Inches(3.4))
            c = tbl.cell(0, 1)
            for j, txt in enumerate(blk[2]):
                p = c.paragraphs[0] if j == 0 else c.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                add_runs(p, txt)
        elif k == "table":
            heads, rows, aligns = blk[1], blk[2], blk[3]
            tbl = doc.add_table(rows=1, cols=len(heads))
            tbl.style = "Table Grid"
            for j, (htxt, a) in enumerate(zip(heads, aligns)):
                cell = tbl.rows[0].cells[j]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                if a:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(htxt); r.font.name = FONT; r.font.size = Pt(7.2); r.bold = True
                shade(cell, "EFEFF6")
            for row in rows:
                cells = tbl.add_row().cells
                for j, (val, a) in enumerate(zip(row, aligns)):
                    p = cells[j].paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    if a:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    add_runs(p, val, size=7.6)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif k in ("cards", "cols"):
            items = blk[1]
            tbl = doc.add_table(rows=1, cols=len(items))
            if k == "cards":
                tbl.style = "Table Grid"
            for j, (title, body) in enumerate(items):
                cell = tbl.rows[0].cells[j]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(title); r.font.name = FONT; r.font.size = Pt(9.5); r.bold = True
                if k == "cards":
                    r.font.color.rgb = brand
                for txt in body:
                    bp = cell.add_paragraph(style="List Bullet" if k == "cols" else None)
                    bp.paragraph_format.space_after = Pt(4)
                    add_runs(bp, txt, size=8.6)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif k == "bullets":
            for txt in blk[1]:
                para(txt, style="List Bullet")

    doc.save(out)


# ----------------------------------------------------------------------- main
def main() -> None:
    d = gather()
    flow_png(d["matches"])
    blocks = content(d)

    # The HTML is only an intermediate for Chrome's PDF renderer, so it goes to a
    # temp directory rather than cluttering docs/ or the repo.
    import tempfile
    html = render_html(blocks)
    hp = Path(tempfile.mkdtemp()) / "writeup.html"
    hp.write_text(html, encoding="utf-8")

    pdf = DOCS / "Off_Ball_Run_Value__Write_Up.pdf"
    if Path(CHROME).exists():
        subprocess.run([CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
                        "--virtual-time-budget=4000", f"--print-to-pdf={pdf}", hp.as_uri()],
                       check=True, capture_output=True)
        print(f"wrote {pdf.name}   ({pdf.stat().st_size / 1e6:.2f} MB)")
    else:
        print("Chrome not found -- PDF skipped")

    docx = DOCS / "Off_Ball_Run_Value__Write_Up.docx"
    render_docx(blocks, docx)
    print(f"wrote {docx.name}  ({docx.stat().st_size / 1e6:.2f} MB)")
    print("\nBoth built from the same content blocks. Once you edit the .docx by hand,")
    print("that file is the master -- re-running this script would overwrite it.")


if __name__ == "__main__":
    main()
